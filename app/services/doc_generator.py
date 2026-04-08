import os
from docx import Document
from app.services.anonymizer import decrypt_name
from datetime import datetime

def create_doc(report_text: str, patient_id: str):
    doc = Document()

    doc.add_heading("Medical Report", 0)
    doc.add_paragraph(f"Patient Name: {decrypt_name(patient_id)}")

    for line in report_text.split("\n"):
        doc.add_paragraph(line)

    file_path = f"{decrypt_name(patient_id)}_{datetime.now().strftime('%Y%m%d')}.docx"
    doc.save(file_path)

    return file_path